import sys, re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')

def add_inline_runs(paragraph, text):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
        else:
            paragraph.add_run(part)

def is_table_sep(line):
    return bool(re.match(r'^\|?[\s:-]+\|[\s:|-]+$', line.strip()))

def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]

def build(md_path, out_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('<!--') or stripped == '':
            i += 1
            continue

        # Table block
        if stripped.startswith('|') and i + 1 < n and is_table_sep(lines[i+1]):
            header = parse_table_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1
            ncols = len(header)
            table = doc.add_table(rows=1, cols=ncols)
            table.style = 'Table Grid'
            for c, htext in enumerate(header):
                cell_p = table.rows[0].cells[c].paragraphs[0]
                add_inline_runs(cell_p, htext)
                for run in cell_p.runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for c in range(ncols):
                    val = row[c] if c < len(row) else ''
                    add_inline_runs(cells[c].paragraphs[0], val)
            doc.add_paragraph('')
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=level)
            add_inline_runs(h, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline_runs(p, text)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped):
            doc.add_paragraph('')
            i += 1
            continue

        # Image marker: ![caption](path)
        m2 = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if m2:
            caption, path = m2.group(1), m2.group(2)
            try:
                doc.add_picture(path, width=Inches(6))
            except Exception:
                p = doc.add_paragraph(f'[Hình: {path} — không nhúng được, chèn thủ công]')
            if caption:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(caption)
                r.italic = True
            i += 1
            continue

        # List item
        m3 = re.match(r'^[-*]\s+(.*)$', stripped)
        if m3:
            p = doc.add_paragraph(style='List Bullet')
            add_inline_runs(p, m3.group(1))
            i += 1
            continue

        m4 = re.match(r'^\d+\.\s+(.*)$', stripped)
        if m4:
            p = doc.add_paragraph(style='List Number')
            add_inline_runs(p, m4.group(1))
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(out_path)
    print('done', out_path)

if __name__ == '__main__':
    build(sys.argv[1], sys.argv[2])
